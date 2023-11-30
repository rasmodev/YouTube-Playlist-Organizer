# Installations
import streamlit as st
from dotenv import dotenv_values
from googleapiclient.discovery import build
from docx import Document

# Load environment variables from the .env file into a dictionary
environment_vars = dotenv_values('.env')

# Get the credential values from the '.env' file
my_api_key = environment_vars.get("API_KEY")

# Enter your API key here
api_key = my_api_key

# Create a YouTube API client
youtube = build('youtube', 'v3', developerKey=api_key)

def main():
    # Page title 
    st.title("YouTube Playlist Video Organizer")
    
    # Add welcome message
    st.markdown("""
    
        Welcome to the YouTube Playlist Video Organizer app! 
        This tool helps you organize videos in a YouTube playlist based on repetitive terms.
    """)

    # Add Banner Image 
    st.image("main.jpg")
    
    # Add sidebar with additional details
    st.sidebar.markdown("""
        **How to Use:**
        - Enter the playlist URL.
        - Specify repetitive terms (comma-separated).
        - Click the "Organize Videos" button.
        - Download the organized videos document.
        Enjoy organizing your YouTube playlist effortlessly!
    """)
        
    # Guide users to replace default values with their own link and repetitive terms
    st.write("_Replace the default values with your own playlist URL and repetitive terms._")
    
    # Get the playlist URL from the user
    playlist_url = st.text_input("**Enter the playlist URL:**", value="https://www.youtube.com/watch?v=rGx1QNdYzvs&list=PLUaB-1hjhk8FE_XZ87vPPSfHqb6OcM0cF")
    
    # Get repetitive terms from the user
    repetitive_terms = st.text_input("**Enter repetitive terms (comma-separated):**", value="Excel, Sql, Power Bi, Tableau, Python")
    repetitive_terms = [term.strip().upper() for term in repetitive_terms.split(',')]

    # Add an organize videos button
    if st.button("**Organize Videos**"):
        organized_doc = organize_videos(playlist_url, repetitive_terms)
        # Provide a download link for the generated document
        st.markdown(get_download_link(organized_doc), unsafe_allow_html=True)
        display_organized_videos(organized_doc)

def organize_videos(playlist_url, repetitive_terms):
    # Create a Word document
    doc = Document()

    # Get the playlist ID from the URL
    playlist_id = playlist_url.split("list=")[1]

    # Define the repetitive terms and headings
    headings = {}

    # Get all the video details from the playlist
    videos = []
    next_page_token = None

    while True:
        # Get the playlist items
        playlist_items = youtube.playlistItems().list(
            part="contentDetails",
            playlistId=playlist_id,
            maxResults=50,
            pageToken=next_page_token
        ).execute()

        # Get the video IDs
        video_ids = [item['contentDetails']['videoId'] for item in playlist_items['items']]

        # Get the video details
        video_details = youtube.videos().list(
            part="snippet",
            id=','.join(video_ids)
        ).execute()

        # Add the videos to the list
        for item in video_details['items']:
            title = item['snippet']['title']
            video_id = item['id']
            video_url = f"https://www.youtube.com/watch?v={video_id}"

            # Group videos based on the repetitive terms
            found = False
            for term in repetitive_terms:
                if term in title.upper():
                    if term not in headings:
                        headings[term] = chr(len(headings) + 65)
                    videos.append({"title": title, "url": video_url, "group": headings[term]})
                    found = True
                    break

            # Add videos without repetitive terms to Introduction or Conclusion
            if not found:
                if item == video_details['items'][0]:
                    videos.append({"title": title, "url": video_url, "group": "Introduction"})
                elif item == video_details['items'][-1]:
                    videos.append({"title": title, "url": video_url, "group": "Conclusion"})

        # Check if there are more pages
        next_page_token = playlist_items.get('nextPageToken')

        if not next_page_token:
            break

    # Print the video titles and URLs grouped by the headings
    for group in sorted(headings.values()):
        doc.add_paragraph(f"\n{group}. {list(headings.keys())[list(headings.values()).index(group)]}")
        count = 1
        for video in videos:
            if video['group'] == group:
                doc.add_paragraph(f"{group}.{count} {video['title']} - {video['url']}")
                count += 1

    # Print videos without repetitive terms under Introduction or Conclusion
    for group in ["Introduction", "Conclusion"]:
        doc.add_paragraph(f"\n{group}")
        count = 1
        for video in videos:
            if video['group'] == group:
                doc.add_paragraph(f"{group}.{count} {video['title']} - {video['url']}")
                count += 1

    return doc

def get_download_link(doc):
    # Save the document to a specific directory
    doc.save("/app/documents/organized_videos.docx")
    # Provide a download link for the document
    doc_download_link = f'<a href="/app/documents/organized_videos.docx" download>Click here to download the organized videos document</a>'
    return doc_download_link

def display_organized_videos(doc):
    st.subheader("Organized Videos")
    for paragraph in doc.paragraphs:
        st.write(paragraph.text)

if __name__ == "__main__":
    main()